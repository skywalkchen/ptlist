import requests
import re
import sys
import os
from docx import Document
from datetime import datetime, timedelta
from docx.shared import Cm
import traceback
from docx.shared import Pt

ptnumberstart=1
location=os.getcwd()
null='0'
session=requests.Session()
atten=''
ward=''
depart=''
consultdep=''
surgery=''
#帳號密碼
while True:
        login_id=input('progressnote登入帳號(DOCXXXXX):')
        try:
                if len(login_id)==5:
                        login_id=re.match(r'\d{5}',login_id)
                        login_id='DOC'+login_id.group(0)
                        break
                else:
                        print('請輸入五碼就好')
        except:
                print('請輸入五碼就好')
while True:
        password=input('progressnote密碼(預設為身份證字號):')
        try:
                if len(password)==10:
                        password=re.match(r'[a-zA-z]\d{9}',password).group(0).upper()
                        break
                else:
                        print('請輸入身份證字號')
        except:
                print('請輸入身份證字號')
#創建patientlist的word檔案
document=Document(location+r'\\default.docx')
section=document.sections[0]
section.left_margin=Cm(0.5)
section.right_margin=Cm(0.5)
section.top_margin=Cm(1.27)
section.bottom_margin=Cm(1.27)

today=datetime.today()
todaydate=str(today.year)+str(today.month).zfill(2)+str(today.day).zfill(2)

paragraph=document.add_paragraph(todaydate+'\n')
#看是要哪一個主治的patient list
atten=input('主治DOC五碼就好，不需要請直接enter:')
if atten!='':
        atten='DOC'+atten
ward=input('請輸入病房號兩碼，不需要請直接enter:')
ward=ward.upper()
depart=input('請輸入科別代號，不需要請直接enter:')
depart=depart.upper()
alldata=input('需要全部的data請輸入1，只要異常值就好請輸入2:')
print('準備中，請稍後......')
#Header
header={
'Connection':'keep-alive',
'Host':'mobilereport.ndmctsgh.edu.tw',
'Origin':'http://mobilereport.ndmctsgh.edu.tw',
'Referer':'http://mobilereport.ndmctsgh.edu.tw/eForm/Account/Login',
'Upgrade-Insecure-Requests':'1',
'User-Agent':'Mozilla/5.0 (Linux; Android 5.0; SM-G900P Build/LRX21T) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/57.0.2987.133 Mobile Safari/537.36'}


#取得verification token
html=session.get('http://mobilereport.ndmctsgh.edu.tw/eForm/Account/Login',headers=header,allow_redirects=False)
cookie_jar=html.cookies
#print (cookie_jar)
#print (html.status_code)
findtoken=re.search(r'RequestVerificationToken" type="hidden" value="(.*)"',html.text)
token=findtoken.group(1)
#print (token)

#登入mobilereport
params={'login_id':login_id,'password':password,'__RequestVerificationToken':token}
s=session.post('http://mobilerep ort.ndmctsgh.edu.tw/eForm/Account/Login',data=params,headers=header,cookies=cookie_jar)
#print (s.status_code)
if s.status_code!=200:
        try:
                sys.exit()
        except:
                print('帳號密碼可能打錯了或是你沒連上mobiletsgh')
#取得交接班清單及病人的資料
s.headers.update({'Referer':'http://mobilereport.ndmctsgh.edu.tw/eForm/Home/Index/'+login_id})
s=session.get('http://mobilereport.ndmctsgh.edu.tw/eForm/PL/ChangeCare/ChangeCareListSearch?SearchChartno=&SearchChinaname=&SearchSectionNo='+depart+'&SearchIDNo=&SearchNRCode='+ward+'&SearchBedCode=&SearchVSDR='+atten+'&SearchSpecialNote=&SearchSKey=&SearchCcGroup=%5B%22%22%5D&SearchDayRegion=1')
if s.status_code!=200:
        print('甚麼東西可能打錯了，請重新執行')
ptlist=re.search(r"patientList = '(.*)'", s.text)
if ptlist==None:
        print('沒有病人，一開始的輸入是對的嗎?請重新啟動')
ptlist=ptlist.group(1).split('},{')
ptlist[0]=ptlist[0][1:]+'}'
for n in range(1,len(ptlist)-1):
        ptlist[n]='{'+ptlist[n]+'}'
ptlist[-1]='{'+ptlist[-1][:-1]
print('共有'+str(len(ptlist))+'位病人')
#取得第n個病人的床號,姓名,姓別,年紀,入院日,病歷號
for n in range(len(ptlist)):
        try:
                ptlist[n]=eval(ptlist[n])
                ptnamegenderage=ptlist[n]['NameGenderAge'].replace(' ','').replace('\n','')
                ptward=ptlist[n]['NrBedNo']
                ptindate=ptlist[n]['INDATETIME'][0:7]
                ptchartno=ptlist[n]['CHARTNO']
                ptchartnoforweb=ptchartno
                ptmedno=ptlist[n]['MEDNO']
                ptvisitseq=ptlist[n]['VISITSEQ']
                if len(ptchartno)==6:
                        ptchartnoforweb='%20'+ptchartno
                pthcaseno=ptlist[n]['HCASENO']
                depname=ptlist[n]['DepName']
                VSname=ptlist[n]['VSDRNAME']
#取得第n個病人的入院impression為active problem
                try:
                        admissionnote=session.get('http://mobilereport.ndmctsgh.edu.tw/mr/AdmissionnoteREPORT.aspx?login_id='+login_id+'&HCASENO='+pthcaseno+'&TITLE='+ptindate+'-'+depname+'&special=n&cno='+ptchartnoforweb)
                        impression=re.search(r"Impression.*\n-*\n(.*)\n-*", admissionnote.text)
                except:
                        print('病人'+ptchartno+'尚未建立admission note')
                        continue
                try:
                        impression=impression.group(1)
                except:
                        impression=''
#TPR
                tprparams={'ptData[CHARTNO]':ptchartno,'ptData[MEDNO]':ptmedno,'ptData[VISITSEQ]':ptvisitseq,'sDate':str((today-timedelta(2)).year)+str((today-timedelta(2)).month).zfill(2)+str((today-timedelta(2)).day).zfill(2),'eDate':str(today.year)+'-'+str(today.month).zfill(2)+'-'+str(today.day).zfill(2)}
                tpr=session.post('http://mobilereport.ndmctsgh.edu.tw/eForm/PL/ChangeCare/VitalSignList',params=tprparams)
                print(tpr.text)
#取得第n個病人已做過data
                today_data='['+str(today.year)+str(today.month).zfill(2)+str(today.day).zfill(2)+']\n'
                one_day_ago_data='['+str((today-timedelta(1)).year)+str((today-timedelta(1)).month).zfill(2)+str((today-timedelta(1)).day).zfill(2)+']\n'
                two_day_ago_data='['+str((today-timedelta(2)).year)+str((today-timedelta(2)).month).zfill(2)+str((today-timedelta(2)).day).zfill(2)+']\n'
                datalist=session.get('http://mobilereport.ndmctsgh.edu.tw/IS3DaysData/Home/ListAllData2?cno='+ptchartnoforweb)
                highdatalist=[]
                lowdatalist=[]
                normaldatalist=[]
                try:
                        highdatalist=re.findall(r'\s*(.*)\s*</td>\s*<td>\s*<span style="color:red;">(.*)</span>\r\n(.*\r\n){16}(.*)',datalist.text)
                except:
                        pass
                for n in range(len(highdatalist)):
                        time=highdatalist[n][3].lstrip()[:10]
                        if highdatalist[n][0]!='Glucose (PC/DEXTRO)\r':
                                if time==str(today.year)+'-'+str(today.month).zfill(2)+'-'+str(today.day).zfill(2):
                                        today_data=today_data+highdatalist[n][0][:-1]+'='+highdatalist[n][1]+'\n'
                                if time==str((today-timedelta(1)).year)+'-'+str((today-timedelta(1)).month).zfill(2)+'-'+str((today-timedelta(1)).day).zfill(2):
                                        one_day_ago_data=one_day_ago_data+highdatalist[n][0][:-1]+'='+highdatalist[n][1]+'\n'
                                if time==str((today-timedelta(2)).year)+str((today-timedelta(2)).month).zfill(2)+str((today-timedelta(2)).day).zfill(2)+'\n':
                                        two_day_ago_data=two_day_ago_data+highdatalist[n][0][:-1]+'='+highdatalist[n][1]+'\n'
                try:
                        lowdatalist=re.findall(r'\s*(.*)\s*</td>\s*<td>\s*<span style="color:orange;">(.*)</span>\r\n(.*\r\n){16}(.*)',datalist.text)
                except:
                        pass
                for n in range(len(lowdatalist)):
                        time=lowdatalist[n][3].lstrip()[:10]
                        if time==str(today.year)+'-'+str(today.month).zfill(2)+'-'+str(today.day).zfill(2):
                                today_data=today_data+lowdatalist[n][0][:-1]+'='+lowdatalist[n][1]+'\n'
                        if time==str((today-timedelta(1)).year)+'-'+str((today-timedelta(1)).month).zfill(2)+'-'+str((today-timedelta(1)).day).zfill(2):
                                one_day_ago_data=one_day_ago_data+lowdatalist[n][0][:-1]+'='+lowdatalist[n][1]+'\n'
                        if time==str((today-timedelta(2)).year)+'-'+str((today-timedelta(2)).month).zfill(2)+'-'+str((today-timedelta(2)).day).zfill(2):
                                two_day_ago_data=two_day_ago_data+lowdatalist[n][0][:-1]+'='+lowdatalist[n][1]+'\n'
                if alldata=='1':
                        try:
                                normaldatalist=re.findall(r'\s*(.*)\s*</td>\s*<td>\s*<span style="color:inherit;">(.*)</span>\r\n(.*\r\n){16}(.*)',datalist.text)
                        except:
                                pass
                        for n in range(len(normaldatalist)):
                                time=normaldatalist[n][3].lstrip()[:10]
                                if time==str(today.year)+'-'+str(today.month).zfill(2)+'-'+str(today.day).zfill(2):
                                        today_data=today_data+normaldatalist[n][0][:-1]+'='+normaldatalist[n][1]+'\n'
                                if time==str((today-timedelta(1)).year)+'-'+str((today-timedelta(1)).month).zfill(2)+'-'+str((today-timedelta(1)).day).zfill(2):
                                        one_day_ago_data=one_day_ago_data+normaldatalist[n][0][:-1]+'='+normaldatalist[n][1]+'\n'
                                if time==str((today-timedelta(2)).year)+str((today-timedelta(2)).month).zfill(2)+str((today-timedelta(2)).day).zfill(2)+'\n':
                                        two_day_ago_data=two_day_ago_data+normaldatalist[n][0][:-1]+'='+normaldatalist[n][1]+'\n'
                else:
                        pass
#取得第n個病人的用藥目前有哪些
                drugweb=session.get('http://mobilereport.ndmctsgh.edu.tw/mr/HISEXNDREPORT.aspx?login_id='+login_id+'&special=n&cno='+ptchartnoforweb)
                try:
                        drugs=re.findall(r'\n\n([\s\S]*)</pre>',drugweb.text)
                        drugs[0]=re.sub('\d\d/\d\d\s\d\d:\d\d','',drugs[0])
                except:
                        drugs=''
#取得最近做的image的日期跟做什麼
                lisreports=''
                lis=session.get('http://mobilereport.ndmctsgh.edu.tw/mr/RisList.aspx?login_id='+login_id+'&special=n&cno='+ptchartnoforweb)
                lisreport=[]
                lisreport=re.findall('>(.*)</a>',lis.text)
                maximage=0
                if len(lisreport)>=4:
                        maximage=4
                else:
                        maximage=len(lisreport)
                for n in range(1,maximage):
                        lisreport[n]='['+lisreport[n][:8]+']'+lisreport[n][9:]
                        lisreport[n]=lisreport[n].replace('PORTABLE CHEST A-P VIEW','CXR')
                        lisreport[n]=lisreport[n].replace('CHEST, P-A VIEW','CXR')
                        lisreport[n]=lisreport[n].replace('CHEST, A-P VIEW','CXR')
                        lisreport[n]=lisreport[n].replace(r'K.U.B. (SUPINE)','KUB')
                        lisreport[n]=lisreport[n].replace('WITHOUT/WITH CONTRAST-C.T.','CT')
                        lisreports=lisreports+lisreport[n]+'\n'
#其它
                others=session.get('http://mobilereport.ndmctsgh.edu.tw/IS3DaysData/Home/ListAllData2?cno='+ptchartnoforweb)
                if re.search('手術排程時間:(.*)',others.text)!=None and re.search('手術處置名稱:(.*)',others.text)!=None:
                        surgery=re.search('手術排程時間:(.*)',others.text).group(1)[:11]+re.search('手術處置名稱:(.*)',others.text).group(1)
                else:
                        surgery='前三日無手術'
                if re.search('會診科別 :(.*)',others.text)!=None:
                        consultdep=re.search('會診科別 :(.*)',others.text).group(0)
                else:
                        consultdep='前三日無會診'
        except Exception:
                print('No.'+str(n)+'病人有問題無法讀取')
                continue
#寫入word      
        locals()['table%s'%n]=document.add_table(rows=8, cols=6)
        ID=locals()['table%s'%n].cell(0,0)
        ID.text='ID  主治：'+VSname
        IDblank=locals()['table%s'%n].cell(1,0).merge(locals()['table%s'%n].cell(4,0))
        IDblank.text=ptnamegenderage+'\n'+ptward+'\n'+ptindate+'\n'+ptchartno
        locals()['table%s'%n].cell(0,1).text='N'
        locals()['table%s'%n].cell(0,2).text='      DIET'
        locals()['table%s'%n].cell(1,1).text='S'
        locals()['table%s'%n].cell(1,2).text='E  M  V  '
        locals()['table%s'%n].cell(2,1).text='V'
        locals()['table%s'%n].cell(2,2).text='SaO2:     ,FiO2:     '
        locals()['table%s'%n].cell(3,1).text='I'
        locals()['table%s'%n].cell(3,2).text='T:    ,ABx   '
        locals()['table%s'%n].cell(4,1).text='P'
        locals()['table%s'%n].cell(4,2).text='P:    ,BP     /    '
        locals()['table%s'%n].cell(0,3).text='Diagnosis'
        diagnosisblank=locals()['table%s'%n].cell(1,3).merge(locals()['table%s'%n].cell(4,3))
        diagnosisblank.text=impression

        Data=locals()['table%s'%n].cell(5,0).merge(locals()['table%s'%n].cell(5,2))
        Data.text='Data'
        Datablank=locals()['table%s'%n].cell(6,0).merge(locals()['table%s'%n].cell(6,2))
        Datablank.text=today_data+one_day_ago_data+two_day_ago_data
        
        Image=locals()['table%s'%n].cell(5,3)
        Image.text='Image'
        Imagetest=locals()['table%s'%n].cell(6,3)
        Imagetest.text=lisreports                                                                                                         
        locals()['table%s'%n].cell(0,4).text='Meds'
        Medsblank=locals()['table%s'%n].cell(1,4).merge(locals()['table%s'%n].cell(6,4))
        Medsblank.text=drugs[0]
        run=Medsblank.paragraphs[0].runs
        run[0].font.size=Pt(8)

        locals()['table%s'%n].cell(0,5).text='Notes'
        Noteblank=locals()['table%s'%n].cell(1,5).merge(locals()['table%s'%n].cell(6,5))
        Noteblank.text=consultdep+surgery
        others=locals()['table%s'%n].cell(7,0).merge(locals()['table%s'%n].cell(7,5))
        others.text='To do:\n\n\n\n'
        print('完成第'+str(ptnumberstart)+'位病人')
        ptnumberstart=ptnumberstart+1
#完成
document.save(location+'\\patientlist'+todaydate+atten+depart+ward+'.docx')
print('大功告成久等啦~上菜')
print('檔案在'+location+'\\patientlist'+todaydate+atten+depart+ward+'.docx')
input()
